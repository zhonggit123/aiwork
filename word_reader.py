# -*- coding: utf-8 -*-
"""
从 Word 文档中读取题目文本，按块切分供 LLM 解析。
"""
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml.ns import qn


def read_word_text(docx_path: str) -> str:
    """读取 Word 全文（仅文字，不含图片）。
    
    特殊处理：若某段落的 run 有下划线格式且文字全为空白，将其转换为 '___'，
    这样下划线空白行（如转述节的填写区域）不会因 strip() 丢失。
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Word 文件不存在: {docx_path}")
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        # 逐 run 处理：下划线+纯空白 → 替换为下划线符号
        run_texts = []
        for run in para.runs:
            if run.underline and run.text and not run.text.strip():
                # 下划线空白 run（学生填写区域）→ 用 ___ 占位
                run_texts.append("_" * max(len(run.text), 8))
            else:
                run_texts.append(run.text)
        text = "".join(run_texts).strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


_OPTION_LABEL_RE = re.compile(r'^[（(]?([A-Da-d])[)）.\s、。，]')


def _convert_to_web_image(image_bytes: bytes) -> tuple:
    """将任意格式图片（如 TIFF/CMYK/调色板）转换为标准 JPEG，返回 (jpeg_bytes, 'jpg')。

    强制转为 RGB 模式，确保浏览器和各平台均可正常渲染。
    若转换失败则尝试输出 PNG，再失败则原样返回（防崩溃）。
    """
    try:
        import io
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        # 统一转为 RGB（处理 CMYK / LAB / 调色板 / RGBA 等所有色彩模式）
        if img.mode == 'RGBA':
            # 透明通道合成白底
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        out = io.BytesIO()
        img.save(out, format='JPEG', quality=85, optimize=True)
        return out.getvalue(), 'jpg'
    except Exception:
        try:
            import io
            from PIL import Image
            img = Image.open(io.BytesIO(image_bytes))
            img = img.convert('RGB')
            out = io.BytesIO()
            img.save(out, format='PNG')
            return out.getvalue(), 'png'
        except Exception:
            return image_bytes, 'jpg'


def _convert_to_png(image_bytes: bytes) -> tuple:
    """将不被浏览器支持的图片格式（如 TIFF）转换为 PNG。
    返回 (bytes, 'png')；若转换失败则返回原始数据和 'png'（扩展名统一）。
    """
    try:
        import io
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        out = io.BytesIO()
        img.save(out, format='PNG')
        return out.getvalue(), 'png'
    except Exception:
        return image_bytes, 'png'


def extract_word_images(docx_path: str) -> list:
    """从 Word 文档中提取内嵌图片（inline shapes），返回图片信息列表。

    每项包含：
      para_index   - 所在段落索引（主体段落顺序，表格内段落为 -1）
      para_text    - 该段落文字（可能为空，例如图片独占一段）
      prev_para_text - 前一段落文字（用于推断选项归属）
      option_label - 推断出的选项字母 A/B/C/D，或 None
      image_bytes  - 图片原始字节
      image_ext    - 扩展名（jpg/png/gif/bmp/webp）
      image_index  - 在本文档中的全局顺序（0-based）
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Word 文件不存在: {docx_path}")

    doc = Document(path)
    results: list = []
    img_idx = 0

    def _extract_from_para(para, para_i: int, prev_text: str) -> None:
        nonlocal img_idx
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if not drawings:
            return
        para_text = para.text.strip()
        _OPT_LETTERS = "ABCD"
        for draw_pos, drawing in enumerate(drawings):
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is None:
                continue
            r_embed = blip.get(qn('r:embed'))
            if not r_embed:
                continue
            try:
                image_part = doc.part.related_parts[r_embed]
                image_bytes = image_part.blob
                content_type = getattr(image_part, 'content_type', 'image/png')
                ext = content_type.split('/')[-1].lower()
                if ext == 'jpeg':
                    ext = 'jpg'
                elif ext not in ('jpg', 'png', 'gif', 'webp'):
                    # TIFF/BMP 等浏览器兼容性差的格式 → 转换为 JPEG
                    image_bytes, ext = _convert_to_web_image(image_bytes)
            except (KeyError, AttributeError):
                continue

            # 推断选项归属：
            # 若同一段落有多张图片（如 A/B/C 都在一行），按段落内顺序 0→A, 1→B, 2→C
            if len(drawings) > 1:
                option_label = _OPT_LETTERS[draw_pos] if draw_pos < len(_OPT_LETTERS) else None
            else:
                # 单图：尝试从本段/前段文字中匹配选项字母
                option_label = None
                for candidate in (para_text, prev_text):
                    m = _OPTION_LABEL_RE.match(candidate)
                    if m:
                        option_label = m.group(1).upper()
                        break

            results.append({
                "para_index": para_i,
                "para_text": para_text,
                "prev_para_text": prev_text,
                "option_label": option_label,
                "image_bytes": image_bytes,
                "image_ext": ext,
                "image_index": img_idx,
            })
            img_idx += 1

    # 遍历主体段落（保留顺序）
    paragraphs = doc.paragraphs
    for para_i, para in enumerate(paragraphs):
        prev_text = paragraphs[para_i - 1].text.strip() if para_i > 0 else ""
        _extract_from_para(para, para_i, prev_text)

    # 遍历表格内段落（表格中的图片）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_paras = cell.paragraphs
                for pi, para in enumerate(cell_paras):
                    prev_text = cell_paras[pi - 1].text.strip() if pi > 0 else ""
                    _extract_from_para(para, -1, prev_text)

    return results


def _render_table_to_image(
    rich_data: list,
    target_size: tuple = (1440, 960),
    padding: int = 40,
    font_size: int = 32,
    merge_info: list = None,
    cols: int = None,
    cell_spans: dict = None,
) -> bytes:
    """将带格式的表格数据渲染为高清 JPEG 图片。

    直接在目标分辨率上绘制，不缩放，保证清晰度。
    默认 1440x960（2倍于 720x480），可根据需要调整。

    Args:
        rich_data: 二维列表，rich_data[row][col] 为 list of {text, bold, underline, newline}
        target_size: 输出图片尺寸 (width, height)
        padding: 表格四周留白
        font_size: 字体大小
        merge_info: 合并信息，merge_info[row][col] = (grid_span, is_continuation)
        cols: 表格的实际列数（网格列数）
        cell_spans: 单元格跨度信息，{(row, col): (grid_span, row_span)}

    Returns:
        JPEG 图片字节
    """
    import io
    from PIL import Image, ImageDraw, ImageFont

    canvas_w, canvas_h = target_size

    if not rich_data or not rich_data[0]:
        img = Image.new('RGB', target_size, (255, 255, 255))
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=95)
        return buf.getvalue()

    rows = len(rich_data)
    if cols is None:
        cols = max(len(row) for row in rich_data)

    # 字体路径（常规 + 粗体）
    font_paths_regular = [
        ("/System/Library/Fonts/Hiragino Sans GB.ttc", 0),  # 冬青黑体 W3（常规）
        ("/System/Library/Fonts/STHeiti Light.ttc", 1),     # Heiti SC Light
        ("C:/Windows/Fonts/msyh.ttc", 0),                   # 微软雅黑
        ("C:/Windows/Fonts/simsun.ttc", 0),
        ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", 0),
    ]
    font_paths_bold = [
        ("/System/Library/Fonts/Hiragino Sans GB.ttc", 2),  # 冬青黑体 W6（粗体）
        ("/System/Library/Fonts/STHeiti Medium.ttc", 1),    # Heiti SC Medium
        ("C:/Windows/Fonts/msyhbd.ttc", 0),                 # 微软雅黑粗体
        ("C:/Windows/Fonts/simhei.ttf", 0),                 # 黑体
        ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", 0),
    ]

    def _load_font(paths, size):
        for fp, idx in paths:
            try:
                return ImageFont.truetype(fp, size, index=idx)
            except Exception:
                continue
        return ImageFont.load_default()

    font_regular = _load_font(font_paths_regular, font_size)
    font_bold = _load_font(font_paths_bold, font_size)

    def _process_runs(cell_runs):
        """处理 runs：下划线空白转为下划线符号，按换行符分行"""
        lines = []
        current_line = []
        for run in cell_runs:
            text = run.get("text", "")
            is_underline = run.get("underline", False)
            is_bold = run.get("bold", False)
            is_newline = run.get("newline", False)

            if is_newline or "\n" in text:
                parts = text.split("\n") if "\n" in text else ["", ""]
                for i, part in enumerate(parts):
                    if i > 0:
                        lines.append(current_line)
                        current_line = []
                    if part:
                        if is_underline and not part.strip():
                            current_line.append({"text": "_" * max(len(part), 6), "bold": False})
                        else:
                            current_line.append({"text": part, "bold": is_bold})
            else:
                if is_underline and not text.strip():
                    current_line.append({"text": "_" * max(len(text), 6), "bold": False})
                else:
                    current_line.append({"text": text, "bold": is_bold})

        if current_line:
            lines.append(current_line)
        return lines

    temp_img = Image.new('RGB', (1, 1))
    temp_draw = ImageDraw.Draw(temp_img)
    line_height = font_size + 4

    col_widths = [0] * cols
    row_heights = [0] * rows

    # 使用 cell_spans 来计算列宽和行高
    for (r_idx, grid_col), (grid_span, row_span) in (cell_spans or {}).items():
        # 找到 rich_data 中对应的单元格内容
        cell_runs = None
        if r_idx < len(rich_data):
            # 计算 rich_data 中的列索引
            data_col = 0
            cur_grid_col = 0
            row_merge = merge_info[r_idx] if merge_info and r_idx < len(merge_info) else None
            for d_idx in range(len(rich_data[r_idx])):
                if row_merge and d_idx < len(row_merge):
                    gs, is_cont = row_merge[d_idx]
                    if is_cont:
                        cur_grid_col += 1
                        continue
                    if cur_grid_col == grid_col:
                        cell_runs = rich_data[r_idx][d_idx]
                        break
                    cur_grid_col += gs if gs > 0 else 1
                else:
                    if cur_grid_col == grid_col:
                        cell_runs = rich_data[r_idx][d_idx]
                        break
                    cur_grid_col += 1
        
        if cell_runs is None:
            cell_runs = []
        
        lines = _process_runs(cell_runs)
        max_line_w = 0
        for line_runs in lines:
            line_text = "".join(r["text"] for r in line_runs)
            try:
                bbox = temp_draw.textbbox((0, 0), line_text, font=font_regular)
                line_w = bbox[2] - bbox[0]
            except Exception:
                line_w = len(line_text) * font_size
            max_line_w = max(max_line_w, line_w)
        num_lines = max(len(lines), 1)
        
        # 列宽：平均分配到跨越的列
        cell_width_per_col = (max_line_w + 16) // grid_span if grid_span > 1 else max_line_w + 16
        for span_idx in range(grid_span):
            if grid_col + span_idx < cols:
                col_widths[grid_col + span_idx] = max(col_widths[grid_col + span_idx], cell_width_per_col)
        
        # 行高：平均分配到跨越的行
        cell_height = num_lines * line_height + 8
        row_height_per_row = cell_height // row_span if row_span > 1 else cell_height
        for span_idx in range(row_span):
            if r_idx + span_idx < rows:
                row_heights[r_idx + span_idx] = max(row_heights[r_idx + span_idx], row_height_per_row)

    col_widths = [max(w, 60) for w in col_widths]
    row_heights = [max(h, 30) for h in row_heights]

    table_w = sum(col_widths) or 1
    table_h = sum(row_heights) or 1

    # 计算缩放比例，让表格撑满画布
    available_w = canvas_w - 2 * padding
    available_h = canvas_h - 2 * padding
    
    # 分别计算宽度和高度的缩放比例
    scale_w = available_w / table_w
    scale_h = available_h / table_h
    
    # 取较大的缩放比例，让表格尽量大（宽度和高度都尽量撑满）
    # 但不能超出画布，所以还是取最小值
    scale = min(scale_w, scale_h)
    
    # 应用缩放到列宽
    scaled_col_widths = [int(w * scale) for w in col_widths]
    scaled_table_w = sum(scaled_col_widths)
    
    # 行高：让高度也撑满，计算需要的额外高度并分配到每行
    base_row_heights = [int(h * scale) for h in row_heights]
    base_table_h = sum(base_row_heights)
    
    # 如果高度还有剩余空间，增加每行高度来撑满
    if base_table_h < available_h:
        extra_h = available_h - base_table_h
        extra_per_row = extra_h // len(base_row_heights)
        scaled_row_heights = [h + extra_per_row for h in base_row_heights]
    else:
        scaled_row_heights = base_row_heights
    
    scaled_table_h = sum(scaled_row_heights)
    
    # 字体大小：按缩放比例，放大 12%
    scaled_font_size = max(int(font_size * scale * 1.12), 20)
    scaled_line_height = scaled_font_size + 6

    font_regular = _load_font(font_paths_regular, scaled_font_size)
    font_bold = _load_font(font_paths_bold, scaled_font_size)

    img = Image.new('RGB', (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    # 表格在画布中居中
    offset_x = (canvas_w - scaled_table_w) // 2
    offset_y = (canvas_h - scaled_table_h) // 2

    # 预计算每行的 y 坐标
    row_y_positions = [offset_y]
    for rh in scaled_row_heights[:-1]:
        row_y_positions.append(row_y_positions[-1] + rh)

    # 遍历每个单元格，使用 cell_spans 来确定跨度
    drawn_cells = set()  # 记录已绘制的单元格，避免重复绘制
    
    for r_idx, row in enumerate(rich_data):
        data_col_idx = 0
        grid_col_idx = 0
        
        row_merge = merge_info[r_idx] if merge_info and r_idx < len(merge_info) else None
        
        while grid_col_idx < cols and data_col_idx < len(row):
            # 检查是否已绘制（被上方单元格纵向合并）
            if (r_idx, grid_col_idx) in drawn_cells:
                # 跳过被合并的单元格
                if row_merge and data_col_idx < len(row_merge):
                    gs, is_cont = row_merge[data_col_idx]
                    if is_cont:
                        data_col_idx += 1
                        continue
                    grid_col_idx += gs if gs > 0 else 1
                else:
                    grid_col_idx += 1
                data_col_idx += 1
                continue
            
            # 获取当前单元格的合并信息
            grid_span = 1
            row_span = 1
            
            if row_merge and data_col_idx < len(row_merge):
                gs, is_cont = row_merge[data_col_idx]
                if is_cont:
                    data_col_idx += 1
                    continue
                grid_span = gs if gs > 0 else 1
            
            # 从 cell_spans 获取纵向跨度
            if cell_spans and (r_idx, grid_col_idx) in cell_spans:
                _, row_span = cell_spans[(r_idx, grid_col_idx)]
            
            # 计算单元格位置和大小
            x = offset_x + sum(scaled_col_widths[:grid_col_idx])
            y = row_y_positions[r_idx]
            cell_w = sum(scaled_col_widths[grid_col_idx:grid_col_idx + grid_span])
            cell_h = sum(scaled_row_heights[r_idx:r_idx + row_span])
            
            # 标记所有被合并的单元格位置
            for ri in range(r_idx, r_idx + row_span):
                for ci in range(grid_col_idx, grid_col_idx + grid_span):
                    drawn_cells.add((ri, ci))
            
            # 获取单元格内容
            cell_runs = row[data_col_idx] if data_col_idx < len(row) else []

            # 绘制单元格边框
            draw.rectangle([x, y, x + cell_w, y + cell_h], outline=(0, 0, 0), width=3)

            if cell_runs:
                lines = _process_runs(cell_runs)
                
                # 单元格内边距
                cell_padding = 6
                max_text_w = cell_w - 2 * cell_padding
                
                # 自动换行
                import re
                wrapped_lines = []
                for line_runs in lines:
                    line_text = "".join(r["text"] for r in line_runs)
                    try:
                        bbox = draw.textbbox((0, 0), line_text, font=font_regular)
                        line_w = bbox[2] - bbox[0]
                    except Exception:
                        line_w = len(line_text) * scaled_font_size
                    
                    if line_w <= max_text_w:
                        wrapped_lines.append(line_runs)
                    else:
                        tokens = []
                        for run in line_runs:
                            run_text = run.get("text", "")
                            is_bold = run.get("bold", False)
                            parts = re.findall(r'[a-zA-Z0-9]+|_+|[^\s\w]+|\s+|[\u4e00-\u9fff]', run_text)
                            for part in parts:
                                tokens.append((part, is_bold))
                        
                        current_runs = []
                        current_w = 0
                        for token_text, is_bold in tokens:
                            fnt = font_bold if is_bold else font_regular
                            try:
                                token_bbox = draw.textbbox((0, 0), token_text, font=fnt)
                                token_w = token_bbox[2] - token_bbox[0]
                            except Exception:
                                token_w = len(token_text) * scaled_font_size
                            
                            if current_w + token_w > max_text_w and current_runs:
                                wrapped_lines.append(current_runs)
                                current_runs = []
                                current_w = 0
                            
                            current_runs.append({"text": token_text, "bold": is_bold})
                            current_w += token_w
                        
                        if current_runs:
                            wrapped_lines.append(current_runs)
                
                num_lines = len(wrapped_lines)
                total_text_h = num_lines * scaled_line_height
                start_y = y + (cell_h - total_text_h) // 2

                for line_idx, line_runs in enumerate(wrapped_lines):
                    line_text = "".join(r["text"] for r in line_runs)
                    try:
                        bbox = draw.textbbox((0, 0), line_text, font=font_regular)
                        line_w = bbox[2] - bbox[0]
                    except Exception:
                        line_w = len(line_text) * scaled_font_size

                    # 第一列居中，其他列左对齐
                    if grid_col_idx == 0:
                        line_x = x + (cell_w - line_w) // 2
                    else:
                        line_x = x + cell_padding
                    line_y = start_y + line_idx * scaled_line_height

                    cursor_x = line_x
                    for run in line_runs:
                        run_text = run.get("text", "")
                        if not run_text:
                            continue
                        is_bold = run.get("bold", False)
                        is_underline = run_text.strip() and all(c == '_' for c in run_text.strip())
                        fnt = font_bold if is_bold else font_regular
                        if is_bold:
                            for dx, dy in [(0, 0), (1, 0), (0, 1), (1, 1)]:
                                draw.text((cursor_x + dx, line_y + dy), run_text, fill=(0, 0, 0), font=fnt)
                        elif is_underline:
                            for dx in [0, 1, 2]:
                                draw.text((cursor_x + dx, line_y), run_text, fill=(0, 0, 0), font=fnt)
                        else:
                            for dx in [0, 1]:
                                draw.text((cursor_x + dx, line_y), run_text, fill=(0, 0, 0), font=fnt)
                        try:
                            run_bbox = draw.textbbox((0, 0), run_text, font=fnt)
                            run_w = run_bbox[2] - run_bbox[0]
                        except Exception:
                            run_w = len(run_text) * scaled_font_size
                        cursor_x += run_w

            grid_col_idx += grid_span
            data_col_idx += 1

    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=95)
    return buf.getvalue()


def extract_word_tables_as_images(
    docx_path: str,
    target_size: tuple = (1440, 960),
) -> list:
    """将 Word 文档中的表格渲染为图片（保留加粗、下划线格式，支持横向和纵向合并单元格）。

    Args:
        docx_path: Word 文件路径
        target_size: 目标图片尺寸 (width, height)

    Returns:
        列表，每项包含：
          table_index  - 表格在文档中的索引（0-based）
          image_bytes  - JPEG 图片字节
          image_ext    - 扩展名 'jpg'
          row_count    - 表格行数
          col_count    - 表格列数
          first_cell   - 第一个单元格内容（用于识别表格类型）
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Word 文件不存在: {docx_path}")

    doc = Document(path)
    results = []

    for t_idx, table in enumerate(doc.tables):
        # 第一遍扫描：收集所有单元格信息，构建合并映射
        # vmerge_starts[grid_col] = (start_row, grid_span) 记录纵向合并的起始位置
        vmerge_starts = {}
        
        # cell_info[row][col] = {grid_span, is_vmerge_start, is_vmerge_continue, cell_runs}
        all_rows_info = []
        
        for row_idx, row in enumerate(table.rows):
            row_info = []
            seen_cells = set()
            
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    # 横向合并的延续部分
                    row_info.append({
                        'grid_span': 0,
                        'is_hmerge_continue': True,
                        'is_vmerge_start': False,
                        'is_vmerge_continue': False,
                        'cell_runs': [],
                    })
                    continue
                seen_cells.add(cell_id)
                
                tc = cell._tc
                grid_span_vals = tc.xpath('.//w:gridSpan/@w:val')
                grid_span = int(grid_span_vals[0]) if grid_span_vals else 1
                
                vmerge_vals = tc.xpath('.//w:vMerge/@w:val')
                vmerge_attr = tc.xpath('.//w:vMerge')
                
                is_vmerge_start = False
                is_vmerge_continue = False
                
                if vmerge_attr:
                    if vmerge_vals and vmerge_vals[0] == 'restart':
                        is_vmerge_start = True
                    else:
                        is_vmerge_continue = True
                
                # 提取单元格内容
                cell_runs = []
                for p_idx, para in enumerate(cell.paragraphs):
                    for run in para.runs:
                        cell_runs.append({
                            "text": run.text,
                            "bold": bool(run.bold),
                            "underline": bool(run.underline),
                            "newline": False,
                        })
                    if p_idx < len(cell.paragraphs) - 1:
                        cell_runs.append({"text": "", "bold": False, "underline": False, "newline": True})
                
                row_info.append({
                    'grid_span': grid_span,
                    'is_hmerge_continue': False,
                    'is_vmerge_start': is_vmerge_start,
                    'is_vmerge_continue': is_vmerge_continue,
                    'cell_runs': cell_runs,
                })
            
            all_rows_info.append(row_info)
        
        if not all_rows_info:
            continue
        
        # 第二遍：计算每个单元格的 row_span
        # cell_spans[(row, grid_col)] = (grid_span, row_span)
        cell_spans = {}
        
        # 计算网格列数
        cols = 0
        for row_info in all_rows_info:
            row_cols = 0
            for ci in row_info:
                if not ci['is_hmerge_continue']:
                    row_cols += ci['grid_span']
            cols = max(cols, row_cols)
        
        # 追踪每个网格列的纵向合并状态
        # vmerge_tracking[grid_col] = (start_row, grid_span)
        vmerge_tracking = {}
        
        # 第一遍：找出所有纵向合并的范围
        # vmerge_ranges[grid_col] = [(start_row, end_row), ...] 每个纵向合并的范围
        vmerge_ranges = {}
        current_vmerge = {}  # grid_col -> start_row
        
        for row_idx, row_info in enumerate(all_rows_info):
            grid_col = 0
            for ci in row_info:
                if ci['is_hmerge_continue']:
                    grid_col += 1
                    continue
                
                grid_span = ci['grid_span']
                
                if ci['is_vmerge_start']:
                    # 开始新的纵向合并
                    current_vmerge[grid_col] = (row_idx, grid_span)
                elif ci['is_vmerge_continue']:
                    # 纵向合并继续
                    pass
                else:
                    # 普通单元格，结束之前的纵向合并（如果有）
                    if grid_col in current_vmerge:
                        start_row, start_span = current_vmerge[grid_col]
                        row_span = row_idx - start_row
                        if row_span >= 1:
                            cell_spans[(start_row, grid_col)] = (start_span, row_span)
                        del current_vmerge[grid_col]
                    # 记录普通单元格
                    cell_spans[(row_idx, grid_col)] = (grid_span, 1)
                
                grid_col += grid_span
        
        # 处理到表格末尾仍在进行的纵向合并
        num_rows = len(all_rows_info)
        for grid_col, (start_row, start_span) in current_vmerge.items():
            row_span = num_rows - start_row
            if row_span >= 1:
                cell_spans[(start_row, grid_col)] = (start_span, row_span)
        
        # 额外检查：如果有单元格没有在 cell_spans 中，但也不是 vmerge_continue，添加它
        for row_idx, row_info in enumerate(all_rows_info):
            grid_col = 0
            for ci in row_info:
                if ci['is_hmerge_continue']:
                    grid_col += 1
                    continue
                grid_span = ci['grid_span']
                if not ci['is_vmerge_continue'] and (row_idx, grid_col) not in cell_spans:
                    cell_spans[(row_idx, grid_col)] = (grid_span, 1)
                grid_col += grid_span
        
        # 构建 rich_data 和 merge_info
        rich_data = []
        merge_info = []
        
        for row_idx, row_info in enumerate(all_rows_info):
            row_data = []
            row_merge = []
            
            for ci in row_info:
                if ci['is_hmerge_continue']:
                    row_merge.append((0, True))
                    continue
                
                if ci['is_vmerge_continue']:
                    row_merge.append((0, True))
                    row_data.append([])
                else:
                    row_merge.append((ci['grid_span'], False))
                    row_data.append(ci['cell_runs'])
            
            rich_data.append(row_data)
            merge_info.append(row_merge)
        
        # 额外处理：检测内容相同的连续单元格（没有使用 vMerge 但内容相同的情况）
        # 这种情况通常出现在 Word 表格被复制/粘贴时
        def get_cell_text(cell_runs):
            return "".join(r.get("text", "") for r in cell_runs).strip()
        
        # 按列检查，找出内容相同的连续单元格
        for grid_col in range(cols):
            # 找到这一列的所有单元格
            col_cells = []  # [(row_idx, data_col_idx, cell_text, grid_span), ...]
            for row_idx, row_info in enumerate(all_rows_info):
                g_col = 0
                d_col = 0
                for ci in row_info:
                    if ci['is_hmerge_continue']:
                        g_col += 1
                        continue
                    if g_col == grid_col:
                        cell_text = get_cell_text(ci['cell_runs'])
                        col_cells.append((row_idx, d_col, cell_text, ci['grid_span'], ci['is_vmerge_continue']))
                        break
                    g_col += ci['grid_span']
                    d_col += 1
            
            # 检查是否有连续相同内容（且内容非空）
            if len(col_cells) >= 2:
                # 找出连续相同内容的范围
                i = 0
                while i < len(col_cells):
                    start_row, start_d_col, start_text, start_span, _ = col_cells[i]
                    if not start_text:  # 跳过空内容
                        i += 1
                        continue
                    
                    # 检查是否已经是 vmerge
                    if (start_row, grid_col) in cell_spans:
                        existing_span = cell_spans[(start_row, grid_col)][1]
                        if existing_span > 1:
                            i += 1
                            continue
                    
                    # 找连续相同的
                    j = i + 1
                    while j < len(col_cells):
                        next_row, next_d_col, next_text, next_span, is_vmerge_cont = col_cells[j]
                        if next_text == start_text and next_span == start_span:
                            j += 1
                        else:
                            break
                    
                    # 如果找到连续相同的（至少2行）
                    if j > i + 1:
                        row_span = j - i
                        # 更新 cell_spans
                        cell_spans[(start_row, grid_col)] = (start_span, row_span)
                        # 更新 merge_info，标记后续行为合并延续
                        for k in range(i + 1, j):
                            merge_row = col_cells[k][0]
                            merge_d_col = col_cells[k][1]
                            if merge_row < len(merge_info) and merge_d_col < len(merge_info[merge_row]):
                                merge_info[merge_row][merge_d_col] = (0, True)
                            # 清空后续行的内容
                            if merge_row < len(rich_data) and merge_d_col < len(rich_data[merge_row]):
                                rich_data[merge_row][merge_d_col] = []
                        i = j
                    else:
                        i += 1
        
        rows = len(rich_data)
        first_cell_text = "".join(r.get("text", "") for r in rich_data[0][0]) if rich_data and rich_data[0] and rich_data[0][0] else ""

        # 渲染为图片，传递合并信息
        image_bytes = _render_table_to_image(
            rich_data, 
            target_size, 
            merge_info=merge_info, 
            cols=cols,
            cell_spans=cell_spans,
        )

        results.append({
            "table_index": t_idx,
            "image_bytes": image_bytes,
            "image_ext": "jpg",
            "row_count": rows,
            "col_count": cols,
            "first_cell": first_cell_text[:50],
        })

    return results


def chunk_text(
    text: str,
    questions_per_batch: int = 5,
    sep: str = "\n\n",
) -> List[str]:
    """
    将长文本按「题」切分成多块。这里用简单策略：按空行分段，每 N 段为一块。
    若你的 Word 里每题格式固定（如都有「题目」「选项」「答案」），可在这里加强逻辑。
    """
    blocks = [b.strip() for b in text.split(sep) if b.strip()]
    if not blocks:
        return []
    chunks = []
    for i in range(0, len(blocks), questions_per_batch):
        chunk = sep.join(blocks[i : i + questions_per_batch])
        chunks.append(chunk)
    return chunks
